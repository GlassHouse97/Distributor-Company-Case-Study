"""Build the polished business-facing Word report for the case study."""

from __future__ import annotations

import csv
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUTS = PROJECT_ROOT / "SOLUTION" / "outputs"
VISUALS = PROJECT_ROOT / "SOLUTION" / "visualizations"
SOLUTION_DIR = PROJECT_ROOT / "SOLUTION"
REPORT_PATH = SOLUTION_DIR / "Distributor_Case_Study_Report.docx"
Q1_PNG = VISUALS / "01_revenue_margin_trends.png"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "20384A"
MUTED = "607184"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"
GRID = "D9E0E7"
GOLD = "B7791F"
WHITE = "FFFFFF"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, size=None, color=INK, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int]):
    if sum(widths_dxa) != 9360:
        raise ValueError(f"Table widths must total 9360 DXA, found {sum(widths_dxa)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
    set_table_borders(table)


def add_table(doc, headers, rows, widths_dxa, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, LIGHT_GRAY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(str(text))
        set_run_font(run, size=9.2, color=NAVY, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row_data):
            cell = cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            if alignments:
                paragraph.alignment = alignments[index]
            run = paragraph.add_run(str(text))
            set_run_font(run, size=9.1, color=INK)
    set_table_geometry(table, widths_dxa)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def add_hyperlink(paragraph, text, url):
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_alt_text(inline_shape, description: str):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def add_picture(doc, path: Path, description: str, width=6.45):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    shape = paragraph.add_run().add_picture(str(path), width=Inches(width))
    set_alt_text(shape, description)
    return shape


def add_caption(doc, text: str):
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    set_run_font(run, size=9, color=MUTED, italic=True)
    return paragraph


def add_callout(doc, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_BLUE)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BLUE)
    borders.append(left)
    p_pr.append(borders)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=NAVY, bold=True)
    return paragraph


def load_csv(name: str):
    with (OUTPUTS / name).open(encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle))


def font(size: int, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/aptos-display.ttf") if bold else Path("C:/Windows/Fonts/aptos.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf") if bold else Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf") if bold else Path("C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def make_q1_chart():
    annual = load_csv("01_annual_revenue_margin.csv")
    width, height = 1600, 900
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = font(30, bold=True)
    subtitle_font = font(18)
    label_font = font(16)
    small_font = font(14)
    bold_font = font(16, bold=True)
    draw.text((70, 48), "Revenue Growth Is Translating Into Stronger Profitability", font=title_font, fill="#17365D")
    draw.text((70, 98), "Recognized revenue composition and gross-margin trend | 2018-2024", font=subtitle_font, fill="#607184")

    left, right, top, bottom = 120, 1480, 205, 625
    max_revenue = 140
    years = [int(row["per_year"]) for row in annual]
    revenue = [float(row["total_revenue"]) / 1_000_000 for row in annual]
    costs = [float(row["total_cost"]) / 1_000_000 for row in annual]
    margins = [float(row["gross_margin_pct"]) for row in annual]
    xs = [left + i * (right - left) / (len(years) - 1) for i in range(len(years))]

    for tick in (0, 35, 70, 105, 140):
        y = bottom - tick / max_revenue * (bottom - top)
        draw.line((left, y, right, y), fill="#D9E0E7", width=2)
        draw.text((left - 25, y), f"${tick}M", font=small_font, fill="#607184", anchor="rm")

    bar_width = 82
    for i, (year, x, rev, cost) in enumerate(zip(years, xs, revenue, costs)):
        total_y = bottom - rev / max_revenue * (bottom - top)
        cost_y = bottom - cost / max_revenue * (bottom - top)
        draw.rectangle((x - bar_width / 2, cost_y, x + bar_width / 2, bottom), fill="#D9EAF7")
        draw.rectangle((x - bar_width / 2, total_y, x + bar_width / 2, cost_y), fill="#2F75B5")
        draw.text((x, bottom + 24), str(year), font=label_font, fill="#20384A", anchor="ma")
        if i in (0, len(years) - 1):
            draw.text((x, total_y - 12), f"${rev:.1f}M", font=bold_font, fill="#17365D", anchor="ms")

    margin_top, margin_bottom = 690, 800
    min_margin, max_margin = 18.5, 22.5
    for tick in (19, 20, 21, 22):
        y = margin_bottom - (tick - min_margin) / (max_margin - min_margin) * (margin_bottom - margin_top)
        draw.line((left, y, right, y), fill="#E7EBF0", width=1)
        draw.text((left - 25, y), f"{tick}%", font=small_font, fill="#607184", anchor="rm")
    points = []
    for x, value in zip(xs, margins):
        y = margin_bottom - (value - min_margin) / (max_margin - min_margin) * (margin_bottom - margin_top)
        points.append((x, y))
    draw.line(points, fill="#D9772A", width=5, joint="curve")
    for i, ((x, y), value) in enumerate(zip(points, margins)):
        draw.ellipse((x - 7, y - 7, x + 7, y + 7), fill="white", outline="#D9772A", width=4)
        if i in (0, len(points) - 1):
            draw.text((x, y - 14), f"{value:.1f}%", font=bold_font, fill="#A64B19", anchor="ms")

    draw.rectangle((1020, 145, 1044, 169), fill="#D9EAF7")
    draw.text((1056, 157), "Cost", font=small_font, fill="#20384A", anchor="lm")
    draw.rectangle((1145, 145, 1169, 169), fill="#2F75B5")
    draw.text((1181, 157), "Gross profit", font=small_font, fill="#20384A", anchor="lm")
    draw.line((1320, 157, 1355, 157), fill="#D9772A", width=5)
    draw.text((1367, 157), "Gross margin", font=small_font, fill="#20384A", anchor="lm")
    draw.text((70, 855), "Complete recognized reporting years; returns and adjustments retained.", font=small_font, fill="#607184")
    image.save(Q1_PNG)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_together = True
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, size=11, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest, size=11)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=11)
    return paragraph


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Caption" not in [style.name for style in doc.styles]:
        doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("REVENUE QUALITY & CUSTOMER PROFITABILITY")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def build_report():
    SOLUTION_DIR.mkdir(exist_ok=True)
    VISUALS.mkdir(exist_ok=True)
    make_q1_chart()

    annual = load_csv("01_annual_revenue_margin.csv")
    segments = load_csv("02_segment_profitability.csv")
    retention = load_csv("04_retention_summary.csv")

    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "Revenue Quality & Customer Profitability Case Study"
    doc.core_properties.subject = "Portfolio case study analyzing growth, profitability, concentration, and retention"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""

    # Editorial cover pattern with generous whitespace and a centered title stack.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(70)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    run = kicker.add_run("PORTFOLIO CASE STUDY")
    set_run_font(run, size=10, color=GOLD, bold=True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Revenue Quality &\nCustomer Profitability")
    set_run_font(run, size=28, color=NAVY, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(34)
    run = subtitle.add_run("A distributor company case study | Recognized-period analysis, 2018-2024")
    set_run_font(run, size=13.5, color=MUTED)
    add_callout(
        doc,
        "Growth was profitable and customer-level diversification was strong. The next value opportunity is to defend the core segment, improve margin where scale is under-monetized, and prioritize retention by recoverable revenue.",
    )
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(38)
    run = note.add_run("Business report | SQL and analyst outputs available in the companion repository")
    set_run_font(run, size=10, color=MUTED, italic=True)

    doc.add_page_break()
    add_heading(doc, "Executive Summary", 1)
    add_body(
        doc,
        "The company’s 2018-2024 growth was profitable and broadly diversified across individual customers. Recognized revenue increased every year, gross profit grew faster, and gross margin expanded. The main strategic exposures are concentrated at the customer-segment level and in a sizable dormant-account pool.",
    )
    add_table(
        doc,
        ["Validated measure", "Result"],
        [
            ["2018-2024 revenue growth", "65.0% ($79.4M to $131.0M)"],
            ["2018-2024 gross-profit growth", "89.6% ($15.2M to $28.8M)"],
            ["Gross-margin change", "+2.86 percentage points"],
            ["Largest segment revenue share", "48.75%"],
            ["Largest customer revenue share", "1.13%"],
            ["Trailing revenue baseline at risk", "$23.0M (17.60% of 2024 revenue)"],
        ],
        [5600, 3760],
    )
    add_heading(doc, "What leadership should do", 2)
    add_table(
        doc,
        ["Priority", "Action", "Expected benefit"],
        [
            ["1", "Protect Independent Retail while expanding high-margin segments", "Defend the main profit pool and improve diversification"],
            ["2", "Launch a monthly value-based retention queue", "Prioritize recent, high-value risk and measure reactivation"],
            ["3", "Review pricing and cost-to-serve in margin-light segments", "Convert existing scale into incremental gross profit"],
        ],
        [900, 4230, 4230],
    )
    add_body(
        doc,
        "The conclusion is descriptive, not predictive. Gross profit excludes operating expenses, inactivity is not confirmed churn, and the revenue-at-risk measure is a trailing historical baseline rather than a forecast of future loss.",
    )

    doc.add_page_break()
    add_heading(doc, "1. Revenue and Margin Trends", 1)
    add_callout(doc, "Recognized revenue rose every year, and gross profit grew faster than revenue.")
    add_picture(
        doc,
        Q1_PNG,
        "Stacked annual revenue bars show cost and gross profit from 2018 through 2024, with an accompanying gross-margin line rising from 19.15 percent to 22.01 percent.",
    )
    add_caption(doc, "Figure 1. Annual recognized revenue composition and gross-margin trend.")
    annual_rows = []
    for row in annual:
        annual_rows.append(
            [
                row["per_year"],
                f"${float(row['total_revenue']) / 1_000_000:.1f}M",
                f"${float(row['gross_profit']) / 1_000_000:.1f}M",
                f"{float(row['gross_margin_pct']):.2f}%",
                "-" if not row["revenue_yoy_pct"] else f"{float(row['revenue_yoy_pct']):.2f}%",
            ]
        )
    add_table(
        doc,
        ["Year", "Revenue", "Gross profit", "Margin", "Revenue YoY"],
        annual_rows,
        [900, 2200, 2200, 1800, 2260],
        [WD_ALIGN_PARAGRAPH.CENTER] + [WD_ALIGN_PARAGRAPH.RIGHT] * 4,
    )
    add_body(
        doc,
        "Revenue increased 65.0% over the period, while gross profit increased 89.6%. Margin expanded in five of six year-over-year comparisons. Revenue growth slowed to about 6% in both 2023 and 2024, but profit still outpaced revenue, which is consistent with maturing rather than deteriorating growth.",
    )

    doc.add_page_break()
    add_heading(doc, "2. Customer Segment Profitability", 1)
    add_callout(doc, "Independent Retail is both the economic engine and the principal segment-level exposure.")
    add_picture(
        doc,
        VISUALS / "02_segment_profitability.png",
        "Horizontal bars compare revenue share and gross-profit share for the ten largest historical customer segments.",
    )
    add_caption(doc, "Figure 2. Segment scale compared with profit contribution.")
    segment_rows = []
    for rank, row in enumerate(segments[:7], start=1):
        segment_rows.append(
            [
                rank,
                row["customer_class_description"].replace("–", "-"),
                f"${float(row['total_revenue']) / 1_000_000:.1f}M",
                f"${float(row['gross_profit']) / 1_000_000:.1f}M",
                f"{float(row['gross_margin_pct']):.2f}%",
                f"{float(row['revenue_share_pct']):.2f}%",
            ]
        )
    add_table(
        doc,
        ["#", "Customer segment", "Revenue", "Gross profit", "Margin", "Rev. share"],
        segment_rows,
        [520, 3040, 1450, 1450, 1450, 1450],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.RIGHT] * 4,
    )
    add_body(
        doc,
        "Independent Retail produced 48.75% of revenue and 51.37% of gross profit at a 22.39% margin. It also generated 43.04% of company revenue growth from 2018 to 2024. Wholesale Trade is more margin-light: 14.12% of revenue but 12.59% of gross profit at an 18.94% margin. Retail Cooperative Network is a credible expansion candidate because its 2024 revenue was 111.18% above 2018 while margin improved to 25.40%.",
    )

    doc.add_page_break()
    add_heading(doc, "3. Revenue Concentration Risk", 1)
    add_callout(doc, "Single-customer dependency is low; segment dependency is materially higher.")
    add_picture(
        doc,
        VISUALS / "03_customer_concentration.png",
        "Cumulative revenue curve shows the share of recognized revenue produced by the ranked customer portfolio, with markers at 50, 80, and 90 percent.",
    )
    add_caption(doc, "Figure 3. Customer-level cumulative revenue concentration.")
    add_table(
        doc,
        ["Concentration measure", "Customer level", "Segment level"],
        [
            ["Largest entity share", "1.13%", "48.75%"],
            ["Top 3 share", "2.97%", "68.82%"],
            ["Top 5 share", "4.31%", "75.91%"],
            ["Entities needed for 80%", "599", "7"],
            ["HHI", "19.36", "2,662.59"],
        ],
        [4160, 2600, 2600],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT],
    )
    add_body(
        doc,
        "The largest customer contributes only 1.13% of revenue, the top 10 contribute 7.09%, and the top 100 contribute 33.41%. It takes 599 customers, or 18.54% of the portfolio, to reach 80% of revenue. Leadership should therefore monitor correlated segment risks—pricing pressure, channel disruption, competitive shifts, or service failures—more closely than the loss of any single account.",
    )

    doc.add_page_break()
    add_heading(doc, "4. Customer Lifecycle and Retention Risk", 1)
    add_callout(doc, "The non-active portfolio carries a $23.0M trailing revenue baseline; 82.27% sits in dormant accounts.")
    add_picture(
        doc,
        VISUALS / "04_retention_risk.png",
        "Stacked bars compare customer lifecycle counts with the distribution of trailing revenue baseline at risk across watch, at-risk, dormant, active, and no-positive-sales groups.",
    )
    add_caption(doc, "Figure 4. Lifecycle mix and trailing revenue baseline at risk.")
    retention_rows = []
    for row in retention:
        retention_rows.append(
            [
                row["churn_risk_bucket"].replace("+", "+"),
                f"{int(row['customer_count']):,}",
                f"{float(row['customer_pct']):.2f}%",
                f"${float(row['revenue_at_risk']) / 1_000_000:.1f}M",
                f"{float(row['revenue_at_risk_share_pct']):.2f}%",
            ]
        )
    add_table(
        doc,
        ["Lifecycle status", "Customers", "Share", "Revenue at risk", "Risk share"],
        retention_rows,
        [3000, 1350, 1350, 1900, 1760],
        [WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.RIGHT] * 4,
    )
    add_body(
        doc,
        "The first outreach wave should focus on high-value customers inactive for 7-12 months, where intervention is still timely. Watch customers should receive lower-cost proactive outreach, while dormant accounts should be handled as a controlled reactivation test. Inactivity is only a behavioral flag; seasonality, planned buying cycles, mergers, or historical class changes may explain the absence of recent sales.",
    )

    doc.add_page_break()
    add_heading(doc, "Strategic Recommendations", 1)
    recommendations = [
        (
            "1. Protect the core while diversifying the growth pipeline",
            "Create an Independent Retail operating review covering revenue, margin, active customers, retention status, service, and company revenue share. In parallel, test repeatable acquisition plays in high-margin segments such as Retail Cooperative Network and Independent Distribution Partner.",
            "Core-segment retention and margin; new revenue and gross profit from targeted expansion segments.",
        ),
        (
            "2. Build a value-based retention queue",
            "Run the lifecycle logic monthly. Prioritize the highest-baseline 7-12 month accounts for human outreach, use low-cost interventions for 4-6 month accounts, and test reactivation only on the most valuable dormant accounts.",
            "Watch-to-active and at-risk-to-active conversion, recovered gross profit, time to intervention, and campaign cost per reactivation.",
        ),
        (
            "3. Improve margin in large, margin-light segments",
            "Start with Wholesale Trade and other material segments whose gross-profit share trails their revenue share. Diagnose price realization, discounting, product mix, purchasing, freight recovery, and service intensity before setting blanket growth targets.",
            "Segment gross margin, gross profit per customer, realized price, freight recovery, returns, and cost-to-serve where available.",
        ),
    ]
    for heading, body, measure in recommendations:
        add_heading(doc, heading, 2)
        add_body(doc, body)
        add_body(doc, f"Primary measures: {measure}", bold_prefix="Primary measures:")
    add_heading(doc, "Management scorecard", 2)
    add_table(
        doc,
        ["Dimension", "Core metric", "Decision signal"],
        [
            ["Growth", "Recognized revenue growth", "Growth slowing without margin support"],
            ["Profitability", "Gross-profit growth and margin", "Revenue-heavy, profit-light segment"],
            ["Concentration", "Largest-segment share", "Segment share rising faster than diversification"],
            ["Retention", "Trailing revenue baseline at risk", "Watch and at-risk pool expanding"],
            ["Seasonality", "Q4 revenue share", "Peak-quarter service or cash strain"],
        ],
        [2100, 3100, 4160],
    )

    doc.add_page_break()
    add_heading(doc, "Methodology, Limitations, and Reproducibility", 1)
    add_heading(doc, "Core analytical rules", 2)
    add_table(
        doc,
        ["Rule", "Application"],
        [
            ["Reporting authority", "Recognized accounting Period in YYMM format; complete window 1801-2412"],
            ["Historical segment", "Customer class recorded on each transaction; current class does not rewrite history"],
            ["Gross profit", "SUM(Sales) - SUM(Cost); aggregate margin equals gross profit divided by revenue"],
            ["Activity definition", "Positive sales reset the lifecycle clock; credits and returns do not"],
            ["Revenue at risk", "Net revenue in the 12 periods ending with last positive sale, floored at zero, for non-active customers"],
            ["Ledger treatment", "Returns, credits, financial edge cases, and duplicate candidates are retained and flagged"],
        ],
        [2700, 6660],
    )
    add_heading(doc, "Important limitations", 2)
    add_body(
        doc,
        "The analysis is descriptive and ends in December 2024. Gross profit excludes operating expenses. Historical segment comparisons include reclassification effects. Inactivity is not confirmed churn, and the revenue-at-risk baseline is not expected future revenue, lifetime value, or forecast loss. Product, market, pricing, service-level, and sales-activity data are not available to establish causal drivers.",
    )
    add_heading(doc, "Reproduce the work", 2)
    add_body(
        doc,
        "The repository contains the final question sheet, all seven anonymized and scaled annual transaction files, portable ingestion scripts, SQL for every analysis, validation checks, compact CSV outputs, and documentation for rebuilding the local SQLite database. The included files can be verified with the published checksum manifest.",
    )
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    prefix = paragraph.add_run("Repository: ")
    set_run_font(prefix, size=11, bold=True)
    add_hyperlink(
        paragraph,
        "Distributor Company Case Study on GitHub",
        "https://github.com/GlassHouse97/Distributor-Company-Case-Study",
    )
    add_body(
        doc,
        "Rebuild sequence: normalize UTF-8 encoding, prepare reference data, build the final SQLite database, run the validation suite, and execute the question-level analysis scripts. The companion Excel workbook provides a compact reviewer path without attempting to place the 3.7-million-row fact table into a worksheet.",
    )
    add_callout(
        doc,
        "Conclusion: The company’s growth history is economically healthy. The highest-value next step is to manage revenue quality explicitly—protecting the profitable core, improving scale economics, and focusing retention effort on the accounts most worth recovering.",
    )

    doc.save(REPORT_PATH)
    print(REPORT_PATH)
    print(Q1_PNG)


if __name__ == "__main__":
    build_report()
